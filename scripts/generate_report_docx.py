"""
Build a Word report from docs/REPORT_SOURCE.md (10-section format).

Edit the Markdown source, then run:
  python scripts/generate_report_docx.py

Optional:
  python scripts/generate_report_docx.py --source path/to/custom.md --out path/to/out.docx
"""
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def strip_html_comments(text: str) -> str:
    return re.sub(r"<!--.*?-->", "", text, flags=re.DOTALL)


def parse_top_level_sections(md: str) -> list[tuple[str, str]]:
    """Split on lines that are exactly '# Title' (one # only)."""
    lines = md.strip().splitlines()
    sections: list[tuple[str, list[str]]] = []
    current_title: str | None = None
    buf: list[str] = []

    for line in lines:
        if re.match(r"^#\s+[^#]", line):
            if current_title is not None:
                sections.append((current_title, buf))
            current_title = line[1:].strip()
            buf = []
        else:
            buf.append(line)
    if current_title is not None:
        sections.append((current_title, buf))
    return [(t, "\n".join(b).strip()) for t, b in sections]


def add_runs_with_bold(paragraph, text: str, font_pt: int = 11) -> None:
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
    for run in paragraph.runs:
        run.font.size = Pt(font_pt)


def add_body_block(doc: Document, block: str, is_references: bool = False) -> None:
    block = block.strip()
    if not block:
        return

    if is_references:
        for line in block.splitlines():
            line = line.strip()
            if re.match(r"^[-*]\s+", line):
                p = doc.add_paragraph(style="List Bullet")
                add_runs_with_bold(p, re.sub(r"^[-*]\s+", "", line))
            elif line:
                p = doc.add_paragraph()
                add_runs_with_bold(p, line)
        return

    if _is_markdown_table(block):
        _add_markdown_table(doc, block)
        return

    parts = re.split(r"\n\n+", block)
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if _is_markdown_table(part):
            _add_markdown_table(doc, part)
        else:
            p = doc.add_paragraph()
            add_runs_with_bold(p, part)


def _is_markdown_table(block: str) -> bool:
    lines = [ln.strip() for ln in block.strip().splitlines() if ln.strip()]
    if len(lines) < 2 or "|" not in lines[0]:
        return False
    sep = lines[1]
    return bool(re.match(r"^\|?[\s\-:|]+\|?", sep))


def _add_markdown_table(doc: Document, block: str) -> None:
    rows_data: list[list[str]] = []
    for line in block.strip().splitlines():
        s = line.strip()
        if not s or "|" not in s:
            continue
        if re.match(r"^\|?[\s\-:|]+\|?\s*$", s):
            continue
        row = s
        if row.startswith("|"):
            row = row[1:]
        if row.endswith("|"):
            row = row[:-1]
        cells = [c.strip() for c in row.split("|")]
        rows_data.append(cells)
    if not rows_data:
        return
    ncols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=ncols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows_data):
        for ci in range(ncols):
            table.rows[ri].cells[ci].text = row[ci] if ci < len(row) else ""
    doc.add_paragraph()


def add_title_page(doc: Document, body: str) -> None:
    parts = [p.strip() for p in re.split(r"\n\n+", body.strip()) if p.strip()]
    if not parts:
        return
    title = parts[0]
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title)
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph()
    for para in parts[1:]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_runs_with_bold(p, para, font_pt=11)


def build_doc(sections: dict[str, str]) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    order = [
        "Title Page",
        "Abstract",
        "Introduction",
        "Literature Review",
        "Methodology",
        "Results",
        "Discussion",
        "Conclusion",
        "References",
        "Appendix (Optional)",
    ]

    # Normalize keys: strip, allow "Appendix" alone
    norm: dict[str, str] = {}
    for k, v in sections.items():
        key = k.strip()
        if key.lower().startswith("appendix"):
            norm["Appendix (Optional)"] = v
        else:
            norm[key] = v

    for i, name in enumerate(order):
        body = norm.get(name)
        if body is None and name == "Appendix (Optional)":
            body = norm.get("Appendix")
        if body is None:
            continue
        if name == "Title Page":
            add_title_page(doc, body)
        else:
            doc.add_heading(name, level=1)
            add_body_block(doc, body, is_references=(name == "References"))
        if i == 0:
            doc.add_page_break()
    return doc


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--source",
        type=Path,
        default=None,
        help="Markdown source (default: docs/REPORT_SOURCE.md)",
    )
    parser.add_argument(
        "--out",
        type=Path,
        default=None,
        help="Output .docx path (default: docs/ESG_Intelligence_Report.docx)",
    )
    args = parser.parse_args()
    base = Path(__file__).resolve().parents[1]
    source = args.source or (base / "docs" / "REPORT_SOURCE.md")
    out_path = args.out or (base / "docs" / "ESG_Intelligence_Report.docx")

    if not source.exists():
        raise SystemExit(f"Source not found: {source}\nCreate it or pass --source.")

    raw = strip_html_comments(source.read_text(encoding="utf-8"))
    pairs = parse_top_level_sections(raw)
    sections = {title: body for title, body in pairs}

    doc = build_doc(sections)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
